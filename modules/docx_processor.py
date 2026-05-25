"""DOCX resume paragraph extraction."""

from __future__ import annotations

from pathlib import Path


def extract_docx_text(file_path: str | Path) -> dict:
    """Extract paragraphs and tables from DOCX."""
    from docx import Document

    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX not found: {path}")

    try:
        doc = Document(str(path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    table_text.append(" | ".join(cells))
        full_text = "\n".join(paragraphs + table_text)
    except Exception as exc:
        raise RuntimeError(f"DOCX extraction failed: {exc}") from exc

    return {
        "text": full_text,
        "metadata": {
            "source": path.name,
            "type": "resume_docx",
            "paragraphs": len(paragraphs),
            "char_count": len(full_text),
        },
    }
